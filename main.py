# 内置模块（优先导入）
import datetime
import os
import re
import getpass

# 第三方核心模块
import yaml
import pytz
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_PARAGRAPH_ALIGNMENT,
    WD_TAB_ALIGNMENT
)
from docx.oxml.ns import qn
from docx.oxml.shared import qn  # 保留（docx不同子模块的qn实际是同一个对象，无需删除）
from docx.shared import Inches, Pt, RGBColor

# 业务相关第三方模块
from uptime_kuma_api import UptimeKumaApi, UptimeKumaException



# --- 基础配置 ---
__version__ = "2.0.0"
CONFIG_FILE = "config.yml"

# --- 启动横幅 ---
def print_banner():
    banner = rf"""
___________                         .___              ________   .________________  ________ 
\_   _____/______   ____   ____   __| _/____   _____  \_____  \  |   ____/   __   \/   __   \
 |    __) \_  __ \_/ __ \_/ __ \ / __ |/  _ \ /     \  /  ____/  |____  \\____    /\____    /
 |     \   |  | \/\  ___/\  ___// /_/ (  <_> )  Y Y  \/       \  /       \  /    /    /    / 
 \___  /   |__|    \___  >\___  >____ |\____/|__|_|  /\_______ \/______  / /____/    /____/  
     \/                \/     \/     \/            \/         \/       \/                    
    Version: {__version__}
    专注生成Word格式Uptime Kuma监控报告
    ====================================================
"""
    print(banner)

# --- 配置管理 ---
def load_config():
    """加载配置文件，仅读取必要的URL、用户名、时区"""
    if not os.path.exists(CONFIG_FILE):
        return None, None, None,None
    try:
        with open(CONFIG_FILE, 'r') as f:
            config = yaml.safe_load(f)
            if config and 'url' in config and 'username' in config:
                url = config['url']
                username = config['username']
                Company= config['Company']
                Company_English_name  = config['Company_English_name']


                print(f"从 {CONFIG_FILE} 加载配置成功。")
                return url, username, Company,Company_English_name
            else:
                print(f"配置文件格式错误，将提示输入新值。")
                return None, None, None, None
    except (yaml.YAMLError, IOError) as e:
        print(f"读取配置文件失败: {e}，将提示输入新值。")
        return None, None, None,None

def save_config(url, username,Company,Company_English_name):
    """保存配置到文件"""
    config_data = {'url': url, 'username': username,'Company': Company , 'Company_English_name' : Company_English_name}
    try:
        with open(CONFIG_FILE, 'w') as f:
            yaml.dump(config_data, f, default_flow_style=False)
        print(f"配置已保存到 {CONFIG_FILE}，下次可直接使用。")
    except IOError as e:
        print(f"保存配置失败: {e}")

def handle_credentials():
    """处理登录凭证（加载配置或手动输入）"""
    url, username, Company, Company_English_name  = load_config()
    save_needed = not all([url, username, Company, Company_English_name])

    if not url:
        url = input("输入Uptime Kuma地址 (例如: http://localhost:3001): ")
    if not username:
        username = input("输入Uptime Kuma用户名: ")
    if not Company:
        Company = input("输入公司名称: ") or '网站监测项目组'
    if not Company_English_name:
        Company_English_name = input("输入公司英文名称，例如（Suzhou Hs Cybersecurity Technology Co., Ltd.): ") or 'Website Monitoring Project Team'
    password = getpass.getpass(f"输入{username}的密码: ")
    return url, username, password,Company, Company_English_name,save_needed

# --- 工具函数 ---


def chose_report():
    # 定义选项：{显示编号: (关键词, 描述)}
    options = {
        1: ("day", "日报（按天统计）"),
        2: ("week", "周报（按周统计）"),
        3: ("month", "月报（按月统计）"),
        4: ("quarter", "季度报告(按季度统计)"),
        5: ("year", "年报(按年统计)"),
    }
    
    # 结构化展示选项
    print("\n📊 请选择报告类型：")
    for num, (key, desc) in options.items():
        print(f"   [{num}] {desc}（可输入关键词：{key}）")
    
    # 构建「输入值→选项」的映射（支持数字/关键词）
    input_mapping = {}
    for num, (key, desc) in options.items():
        input_mapping[str(num)] = (num, key, desc)  # 数字字符串
        input_mapping[key.lower()] = (num, key, desc)  # 关键词
    
    # 循环直到输入有效
    while True:
        user_input = input("\n请输入编号/关键词（默认3=月报）：").strip().lower()
        # 处理默认选择
        if not user_input:
            user_input = "3"
        
        # 校验并返回
        if user_input in input_mapping:
            num, key, desc = input_mapping[user_input]
            print(f"✅ 已选择：{key}")
            return key
        else:
            # 友好提示：列出有效输入
            valid_inputs = list(input_mapping.keys())
            print(f"❌ 无效输入！有效选项：{valid_inputs}，请重新输入。")

def calculate_hours_since_period_start(period) :
    """
    计算指定时间维度的第一天距离当前时间的小时数（保留2位小数）
    
    :return: 距离当前时间的小时数（正数，保留2位小数）
    :raises ValueError: 无效的period/时区/周起始日
    :raises Exception: 其他时间计算异常
    """
    
    
    
    # 1. 校验入参合法性
    valid_periods = ["day", "week", "month", "quarter", "year"]
    if period.lower() not in valid_periods:
        raise ValueError(f"无效的period！仅支持：{', '.join(valid_periods)}")
    
    
    
    try:
        # 2. 初始化时区对象
        timezone = "Asia/Shanghai"
        tz = pytz.timezone(timezone)
    except pytz.UnknownTimeZoneError:
        raise ValueError(f"无效的时区：{timezone}，可选值参考pytz.all_timezones")
    
    # 3. 获取当前带时区的时间（精确到秒）
    now = datetime.datetime.now(tz).replace(second=0, microsecond=0)
    
    # 4. 根据不同period计算「第一天」的0时0分0秒
    period_start = None

    
    if period.lower() == "day":
        # 本日：当前日期的0时0分0秒
        period_start = now.replace(hour=0, minute=0)

    
    elif period.lower() == "week":
        # 本周：周起始日的0时0分0秒（默认周一）
        current_weekday = now.weekday()  # 0=周一，6=周日
        # 计算距离周起始日的天数差
        day_diff = (current_weekday ) % 7
        period_start = (now - datetime.timedelta(days=day_diff)).replace(hour=0, minute=0)

    
    elif period.lower() == "month":
        # 本月：当月1号的0时0分0秒
        period_start = now.replace(day=1, hour=0, minute=0)

    
    elif period.lower() == "quarter":
        # 本季度：季度第一天（1/4/7/10月1号）的0时0分0秒
        # 计算当前季度的起始月份：(当前月-1)//3 *3 +1
        quarter_start_month = ((now.month - 1) // 3) * 3 + 1
        period_start = now.replace(month=quarter_start_month, day=1, hour=0, minute=0)

    
    elif period.lower() == "year":
        # 本年：1月1号的0时0分0秒
        period_start = now.replace(month=1, day=1, hour=0, minute=0)

    
    # 5. 计算时间差并转换为小时数（不保留小数）
    if period_start != None:
        time_diff = now - period_start
        hours_diff = round(time_diff.total_seconds() / 3600, )  
    return hours_diff


# --- 数据处理 ---
def analyze_heartbeats(heartbeats):
    """分析心跳数据，计算停机事件和延迟数据"""
  
    user_tz = pytz.utc

    def to_datetime(time_val):
        """转换时间为带时区的datetime对象"""
        if isinstance(time_val, str):
            try:
                naive_dt = datetime.datetime.strptime(time_val.split('.')[0], '%Y-%m-%d %H:%M:%S')
                return pytz.utc.localize(naive_dt).astimezone(user_tz)
            except ValueError:
                return None
        elif isinstance(time_val, (int, float)):
            return datetime.datetime.fromtimestamp(float(time_val), tz=pytz.utc).astimezone(user_tz)
        return None

    processed_beats = []
    ping_data = []
    keywords = []
    keywords_count=0
    count=0


    # print(heartbeats[0])
    for beat in heartbeats:
        count+=1
        msg = beat['msg']
        if 'but' in msg:
           
            pattern = r"\[(.*?)\]"
            match = re.search(pattern, msg)
            if match:
                
                results = match.group(1)
                if results != "":
                    for i in results:
                        keywords.append(results)
                    keywords_count+=1
                    
                    
        dt = to_datetime(beat.get('time'))
        if dt:
            processed_beats.append({'datetime': dt, 'status': beat['status']})
            if beat.get('ping') is not None:
                ping_data.append({'datetime': dt, 'ping': beat['ping']})
    # print(list(set(keywords)))
    # print(keywords_count)
    #  计算关键词占比
    unique_keywords = list(set(keywords))
    keyword_ratio = (keywords_count / count * 100) if count > 0 else 0.0
    # 统计每个关键词出现次数
   
    beats = sorted(processed_beats, key=lambda x: x['datetime'])
    incidents = []
    current_downtime_start_dt = None

    keyword_analysis = {
        "unique_keywords": unique_keywords,
        "keyword_count": keywords_count,
        "keyword_ratio": round(keyword_ratio, 2)   
    }

    # 识别停机事件
    for beat in beats:
        is_down = beat['status'] == 0
        if is_down and current_downtime_start_dt is None:
            current_downtime_start_dt = beat['datetime']
        elif not is_down and current_downtime_start_dt is not None:
            incidents.append({
                "start": current_downtime_start_dt,
                "duration": beat['datetime'] - current_downtime_start_dt
            })
            current_downtime_start_dt = None

    # 处理持续中的停机
    if current_downtime_start_dt is not None:
        now_aware = datetime.datetime.now(user_tz)
        incidents.append({"start": current_downtime_start_dt, "duration": now_aware - current_downtime_start_dt, "ongoing": True})
    # print(keyword_analysis)
    return {"downtime_incidents": incidents,"keyword_analysis":keyword_analysis, "ping_data": ping_data}

def calculate_summary_stats(analysis_results):
    """计算日/周/月维度的汇总统计"""
    incidents = analysis_results['downtime_incidents']
    ping_data = analysis_results['ping_data']

    
    user_tz = pytz.utc

    now = datetime.datetime.now(user_tz)
    periods = {
        "日": datetime.timedelta(days=1),
        "周": datetime.timedelta(days=7),
        "月": datetime.timedelta(days=30)
    }

    summary = {}
    for name, delta in periods.items():
        period_start = now - delta

        # 停机统计
        period_incidents = [inc for inc in incidents if inc['start'] >= period_start]
        count = len(period_incidents)
        total_duration = sum([inc['duration'] for inc in period_incidents], datetime.timedelta())
        avg_duration = total_duration / count if count > 0 else datetime.timedelta(0)
        percentage = (total_duration.total_seconds() / delta.total_seconds()) * 100 if delta.total_seconds() > 0 else 0

        # 延迟统计
        period_pings = [p['ping'] for p in ping_data if p['datetime'] >= period_start]
        avg_ping = sum(period_pings) / len(period_pings) if period_pings else None
        max_ping = max(period_pings) if period_pings else None

        summary[name] = {
            "count": count,
            "avg_duration": avg_duration,
            "percentage": percentage,
            "avg_ping": avg_ping,
            "max_ping": max_ping
        }

    return summary


# --- 监控项选择 ---
def select_monitors(monitors):
    """
    让用户选择需要生成报告的监控项（仅展示parent为None的顶级监控项，区分监控组/单独监控）
    
    :param monitors: 监控项列表（从Uptime Kuma API获取的原始列表）
    :return: 选中的监控项列表；若无有效监控项/用户取消，返回空列表
    """

    valid_display_ids = []  # 存储有效的显示编号（用户看到的ID）
    display_mapping = {} 

    print("\n📋 可用监控项:")
    for monitor in monitors:
        # 仅处理parent为None的顶级监控项
        if monitor.get('parent') is None:
            # 区分监控组和单独监控
            pathName=monitor["pathName"]
            ID = monitor["id"]
            if monitor.get("childrenIDs", []) != []:   
                print(f"ID: {ID},监控组: {pathName}")
                display_mapping[ID] = monitor["childrenIDs"]
            else:
                print(f"ID: {ID},单独监控: {pathName}")
                display_mapping[ID] = ID
            valid_display_ids.append(ID)

    if valid_display_ids is None:
        for monitor in monitors:
            pathName=monitor["pathName"]
            ID = monitor["id"]
            print(f"ID: {ID},子监控: {pathName}")
            valid_display_ids.append(ID)
            display_mapping[ID] = ID
    # print(display_mapping)
    # 4. 用户输入选择（循环直到输入有效）
    prompt = f"\n请输入监控项编号（有效编号：{valid_display_ids}），多个编号用逗号分隔,监控组建议只选一个："
    
    while True:
        try:
            selection = input(prompt).strip()
            
            # 处理空输入
            if not selection:
                print("❌ 输入不能为空，请重新输入！")
                continue
            # 解析用户输入的编号（去重、排序）
            selected_display_ids = [int(s.strip()) for s in selection.split(',')]
            selected_display_ids = list(set(selected_display_ids))  # 去重
            selected_display_ids.sort()  # 排序
            
            # 验证编号是否有效
            invalid_ids = [sid for sid in selected_display_ids if sid not in valid_display_ids]
            if invalid_ids:
                print(f"❌ 无效编号：{invalid_ids}，有效编号范围：{valid_display_ids}，请重新输入！")
                continue
            # 初始化空列表，用于存储选中项对应的原监控项列表索引
            selected_list_indices = []

            # 遍历用户选中的显示编号（已去重+排序）
            for sid in selected_display_ids:
                # 从映射字典中获取该显示编号对应的原列表索引
                original_index = display_mapping[sid]
                if type(original_index) == list:
                    selected_list_indices.append(sid)
                    for i in original_index:
                        selected_list_indices.append(i)
                else:
                    # 将索引追加到列表中
                    selected_list_indices.append(original_index)
            return selected_list_indices

            
        
        except ValueError:
            print("❌ 输入格式错误！请输入数字（多个用逗号分隔）")
        except Exception as e:
            print(f"❌ 输入处理出错：{str(e)}，请重新输入！")


# --- Word报告生成核心函数 ---
def generate_docx_report(project_name,period,Company, Company_English_name, selected_monitors, all_monitor_data):
    """生成Word格式的监控报告"""
    # 创建Word文档
    if project_name is None:
        doc_name='\n网站'
    else:
        doc_name=project_name+"\n网站"


    if project_name is not None:

        # 分隔信息
        project_target = None
        project_monitor_data =None
        for m in selected_monitors:
            if m['name'] == project_name:
                project_target = m
                break

        if project_target:
            selected_monitors.remove(project_target)

        for d in all_monitor_data:
            if d['monitor_name'] == project_name:
                project_monitor_data = d
                break

        if project_monitor_data:
            all_monitor_data.remove(project_monitor_data)


    url_list = []
    for monitor in selected_monitors:
        url_dic= {
            'name': monitor["name"],
            'url': monitor["url"]
                   }
        url_list.append(url_dic)

    doc = Document()
    section = doc.sections[0]
    header = section.header
    # 清除页眉默认空段落（避免多余空行）
    for para in header.paragraphs:
        para.clear()
    # 3. 基础页眉设置（所有页面共用）
    header_para = header.add_paragraph()
    header_para.paragraph_format.first_line_indent = Pt(0)
    header_para.paragraph_format.line_spacing=Pt(0)
    # header_para.paragraph_format.line_spacing = Pt(20)
    # 1. 添加左侧图片（关键：设置垂直对齐为居中）
    run_img = header_para.add_run()
    run_img.add_picture("./logo/logo.png", width=Inches(1))


    # 2. 添加制表符实现文字右对齐
    if section:
        tab_stop = header_para.paragraph_format.tab_stops.add_tab_stop(
            section.page_width - section.left_margin - section.right_margin, # pyright: ignore[reportOperatorIssue]
            WD_TAB_ALIGNMENT.RIGHT
        )
    header_para.add_run('\t')  # 插入制表符跳转到右侧

    # 3. 添加右侧文字
    run_text = header_para.add_run("网站监测服务报告")
    run_text.font.name = '宋体'
    run_text._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_text.font.size = Pt(14)
    run_text.font.bold = True

    valid_periods = {"day":"日报", "week":"周报", "month":"月报", "quarter":"季度报告", "year":"年报"}
    report_period=valid_periods[period]
    # print(report_period)
    # 设置全局字体（兼容中英文）
    normal_style =  doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_para_format = normal_style.paragraph_format

    # 2.1 设置全局1.5倍行距
    normal_para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 2.2 可选：同时设置全局首行缩进、段间距（按需添加）
    normal_para_format.first_line_indent = Pt(24)  # 全局首行缩进2字符
    normal_para_format.space_before = Pt(0)        # 全局段前间距5磅
    normal_para_format.space_after = Pt(0)         # 全局段后间距5磅

    Heading_1_style =  doc.styles['Heading 1']
    Heading_1_style.font.name = '宋体'
    Heading_1_style.font.size = Pt(16)
    Heading_1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Heading_1_style.font.color.rgb = RGBColor(0, 0, 0)
    Heading_1_style_format = Heading_1_style.paragraph_format
    Heading_1_style_format.space_before = Pt(0)
    Heading_1_style_format.space_after = Pt(0)


    Heading_2_style =  doc.styles['Heading 2']
    Heading_2_style.font.name = '宋体'
    Heading_2_style.font.size = Pt(15)
    Heading_2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Heading_2_style.font.color.rgb = RGBColor(0, 0, 0)
    Heading_2_style_format = Heading_2_style.paragraph_format
    Heading_2_style_format.space_before = Pt(0)        # 全局段前间距0磅
    Heading_2_style_format.space_after = Pt(0)         # 全局段后间距0磅

    Heading_3_style =  doc.styles['Heading 3']
    Heading_3_style.font.name = '宋体'
    Heading_3_style.font.size = Pt(14)
    Heading_3_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Heading_3_style.font.color.rgb = RGBColor(0, 0, 0)
    Heading_3_style_format = Heading_3_style.paragraph_format
    Heading_3_style_format.space_before = Pt(0)        # 全局段前间距0磅
    Heading_3_style_format.space_after = Pt(0)         # 全局段后间距0磅


    Heading_4_style =  doc.styles['Heading 4']
    Heading_4_style.font.name = '宋体'
    Heading_4_style.font.size = Pt(13)
    Heading_4_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    Heading_4_style.font.color.rgb = RGBColor(0, 0, 0)
    Heading_4_style_format = Heading_4_style.paragraph_format
    Heading_4_style_format.space_before = Pt(0)        # 全局段前间距0磅
    Heading_4_style_format.space_after = Pt(0)
    Heading_4_style.font.italic = False


    doc.add_paragraph("\n\n\n\n\n\n\n\n")

    # 添加报告标题
    title = doc.add_heading(f'{doc_name}检测服务{report_period}', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = '黑体'
    title.paragraph_format.first_line_indent = Pt(0)
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph("\n\n\n\n\n\n\n")



    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 表格在页面行内居中

    # 2. 获取单元格并设置内容+格式
    cell = table.rows[0].cells[0]

    for para in cell.paragraphs:
        para.clear()

    para = cell.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(0)
    
    run1 = para.add_run(Company+'\n'+Company_English_name)


    run1.font.name = '黑体'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run1.font.size = Pt(12)
    run1.font.bold = True

    # 5. 设置单元格内文本分散对齐（关键：段落水平分散对齐 + 单元格垂直居中）
    # 5.1 段落水平分散对齐（文字左右均匀分布）
    para.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE  # 分散对齐（兼容中文）
    # 5.2 单元格垂直居中（文本在单元格内上下居中）
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 6. 可选：调整单元格宽度（让分散对齐效果更明显）
    cell.width = Pt(340)  # 设置单元格宽度为400磅，便于分散对齐展示


    # 添加基础信息
    now_aware = datetime.datetime.now(pytz.timezone("UTC"))
    generated_str = now_aware.strftime('%Y-%m-%d')
    date = doc.add_paragraph(f"\n{generated_str}")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.first_line_indent = Pt(0)









    doc.add_page_break()










    title1 = doc.add_heading('一、 综述信息', level=1)
    title1_run = title1.runs[0]
    title1_run.font.name = '宋体'
    title1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title2 = doc.add_heading('1. 监测概述', level=2)
    title2_run = title2.runs[0]
    title2_run.font.name = '宋体'
    title2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph(f'为持续保障客户核心互联网资产的稳定运行、合规发布及信息安全，{Company}（以下简称“我方”）针对性部署了多维度网站安全监测系统，构建“实时监测-智能告警-人工核查-快速处置”的全流程主动防御体系，实现7×24小时不间断监测覆盖，最大限度降低安全风险及业务中断损失。')



    # 监测概述正文第二段
    doc.add_paragraph('本周期内，监测系统围绕客户指定网站资产，聚焦核心安全及性能指标开展全方位监测，包括但不限于：')


    n1 = doc.add_paragraph(style='List Bullet')
    n1.paragraph_format.first_line_indent = Pt(24)
    n1.add_run('可用性监测：').bold = True
    n1.add_run('采用定时主动探测机制，按预设周期对目标站点发起标准化访问请求，全面校验服务连通性、响应状态、页面加载时效与跳转逻辑，精准识别无法访问、连接超时、异常跳转、服务中断等可用性风险，确保业务链路持续稳定可用；')
    n2 = doc.add_paragraph(style='List Bullet')
    n2.paragraph_format.first_line_indent = Pt(24)
    n2.add_run('内容合规监测：').bold = True
    n2.add_run('于实时内容巡检与智能识别能力，对页面文本、元素及关键信息进行全量扫描核验，自动排查违法违规内容、敏感信息、不当表述及不合规要素，实现风险内容早发现、早预警，保障平台内容安全与合规运营；')



    # 监测概述最后一段
    doc.add_paragraph('监测过程中，系统一旦捕获上述异常指标，将立即触发分级告警机制，通过专属邮件通道实时推送至指定监测工程师。工程师在收到告警后30分钟内启动人工核查，结合业务场景开展风险研判，同步形成初步处置建议，协助客户快速响应、闭环处置，最大限度控制安全事件影响范围及损失。')


    title3= doc.add_heading('2. 监测对象', level=2)
    title3_run = title3.runs[0]
    title3_run.font.name = '宋体'
    title3_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # doc.add_paragraph(f"监控项: {monitor_names}")
    # 创建监测对象表格：2列（系统名称、访问地址），首行为表头
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    # table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格左对齐
    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '系统名称'
    hdr_cells[1].text = '访问地址'
    # 格式化表头字体（宋体10号、加粗）
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
    # 动态添加监测对象数据
    for urlinfo in url_list:
        row_cells = table.add_row().cells
        row_cells[0].text = urlinfo.get('name', '')
        row_cells[1].text = urlinfo.get('url', '')
        # 格式化单元格内容（宋体10号、常规）
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)


    note_title = doc.add_paragraph()
    note_title.add_run('注：').bold = True
    note_title.add_run('本系统在实施监测过程中，受限于以下客观环境因素，可能导致部分监测覆盖度受到影响：')
    note_title.paragraph_format.first_line_indent = Pt(24)  # 取消首行缩进

    n3 = doc.add_paragraph(style='List Bullet')
    n3.paragraph_format.first_line_indent = Pt(24)
    n3.add_run('安全设备拦截限制：').bold = True
    n3.add_run('目标网站部署的防护机制（如WAF、防火墙等）可能将系统高频、深度的探测行为识别为恶意攻击，进而触发拦截机制，导致影响监测全面性；')
    n4 = doc.add_paragraph(style='List Bullet')
    n4.paragraph_format.first_line_indent = Pt(24)
    n4.add_run('认证页面访问受限：').bold = True
    n4.add_run('由于未配置登录凭证，系统无法进入需身份验证的后台或受保护区域，故对登录后的功能模块、动态内容及深层业务逻辑暂无法开展监测评估。')





    title4 = doc.add_heading('二、监测结果', level=1)
    title4_run = title4.runs[0]
    title4_run.font.name = '宋体'
    title4_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 监控组信息详情
    ti = 1
    if project_name is not None:
        # 添加项目总览
        if project_monitor_data:
            monitor_name = project_monitor_data['monitor_name']
            summary_stats = project_monitor_data['summary_stats']
            title5 = doc.add_heading(f"1. 监控项目总览: {monitor_name}", level=2)
            title5_run = title5.runs[0]
            title5_run.font.name = '宋体'
            title5_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            ti = 2


        # 汇总统计表格
        summary_table = doc.add_table(rows=1, cols=6)
        summary_table.style = 'Table Grid'
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = '统计维度'
        hdr_cells[1].text = '停机次数'
        hdr_cells[2].text = '平均停机时长'
        hdr_cells[3].text = '平均延迟'
        hdr_cells[4].text = '最大延迟'
        hdr_cells[5].text = '停机占比'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 填充统计数据
        for period, stats in summary_stats.items():
            row_cells = summary_table.add_row().cells
            row_cells[0].text = period
            row_cells[1].text = str(stats['count'])
            row_cells[2].text = _format_timedelta(stats['avg_duration'])
            row_cells[3].text = f"{int(stats['avg_ping'])} ms" if stats['avg_ping'] else "N/A"
            row_cells[4].text = f"{int(stats['max_ping'])} ms" if stats['max_ping'] else "N/A"
            row_cells[5].text = f"{stats['percentage']:.2f}%"
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10)
                        run.font.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        title5 = doc.add_heading(f"1. 监控详情", level=2)
        title5_run = title5.runs[0]
        title5_run.font.name = '宋体'
        title5_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')



    # 为每个监控项添加详情
    for idx, data in enumerate(all_monitor_data):
        # print(data)
        monitor_name = data['monitor_name']
        summary_stats = data['summary_stats']
        incidents = data['downtime_incidents']
        keyword_analysis = data['keyword_analysis']


        title6 = doc.add_heading(f"{ti}.{idx+1}. 监控项: {monitor_name}", level=3)
        title6_run = title6.runs[0]
        title6_run.font.name = '宋体'
        title6_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

       # 汇总统计表格
        summary_table = doc.add_table(rows=1, cols=6)
        summary_table.style = 'Table Grid'
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = '统计维度'
        hdr_cells[1].text = '停机次数'
        hdr_cells[2].text = '平均停机时长'
        hdr_cells[3].text = '平均延迟'
        hdr_cells[4].text = '最大延迟'
        hdr_cells[5].text = '停机占比'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)

        # 填充统计数据
        for period, stats in summary_stats.items():
            row_cells = summary_table.add_row().cells
            row_cells[0].text = period
            row_cells[1].text = str(stats['count'])
            row_cells[2].text = _format_timedelta(stats['avg_duration'])
            row_cells[3].text = f"{int(stats['avg_ping'])} ms" if stats['avg_ping'] else "N/A"
            row_cells[4].text = f"{int(stats['max_ping'])} ms" if stats['max_ping'] else "N/A"
            row_cells[5].text = f"{stats['percentage']:.2f}%"
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10)
                        run.font.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)

        # 关键词事件日志
        title7 = doc.add_heading(f"{ti}.{idx+1}.1. 关键词日志", level=4)
        title7_run = title7.runs[0]
        title7_run.font.name = '宋体'
        title7_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if keyword_analysis["keyword_count"] == 0:
            doc.add_paragraph("该时间段内无关键词事件")
        else:
            keywords = str(keyword_analysis['unique_keywords']) or "N/A"
            keyword_count =str(keyword_analysis['keyword_count']) or "N/A"
            keyword_ratio =str(keyword_analysis['keyword_ratio']) or "N/A"
            doc.add_paragraph(f"累计触发以下关键词:", style='List Bullet')
            doc.add_paragraph(f"{keywords}", style='List Bullet')
            doc.add_paragraph(f"总计占比: {keyword_ratio}，共{keyword_count}次", style='List Bullet')
            doc.add_paragraph()  # 空行分隔

        # 停机事件日志
        title8 = doc.add_heading(f"{ti}.{idx+1}.2. 停机事件日志（时间排序）", level=4)
        title8_run = title8.runs[0]
        title8_run.font.name = '宋体'
        title8_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if not incidents:
            doc.add_paragraph("该时间段内无停机事件")
        else:
            for incident in reversed(incidents):
                start_str = incident['start'].strftime('%Y-%m-%d %H:%M:%S %Z')
                duration_str = _format_timedelta(incident['duration'])
                if incident.get("ongoing", False):
                    duration_str += " (持续中)"

                doc.add_paragraph(f"停机开始: {start_str}", style='List Bullet').paragraph_format.first_line_indent = Pt(24)
                doc.add_paragraph(f"持续时长: {duration_str}", style='List Bullet').paragraph_format.first_line_indent = Pt(24)
                # doc.add_paragraph()  # 空行分隔


    # 结果总结和进一步规划
    title9 = doc.add_heading('三、监控结果总结', level=1)
    title9_run = title9.runs[0]
    title9_run.font.name = '宋体'
    title9_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph("本次监控周期内，系统围绕目标站点的可用性、内容合规等核心维度，开展常态化、全覆盖、自动化监控工作，全面排查站点运行过程中的可用性风险与内容合规隐患，确保站点稳定、合规运营，现将监控结果、通用修复建议及下一步监测计划总结如下：")

    title10 = doc.add_heading('1. 监控结果概述', level=2)
    title10_run = title10.runs[0]
    title10_run.font.name = '宋体'
    title10_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph("本次监控覆盖站点全业务页面及核心访问链路，监测过程规范、数据精准，整体运行情况如下：")

    n6 = doc.add_paragraph(style='List Bullet')
    n6.add_run('可用性监测：').bold = True
    n6.add_run('监控周期内，按预设固定周期发起标准化访问请求，全面核查站点连通性、响应时效、页面加载状态及跳转逻辑，重点排查无法访问、连接超时、异常跳转、服务中断等典型故障。经全面监测，站点整体访问稳定性良好，核心业务链路响应正常，未出现重大可用性故障；若存在零星轻微异常（如瞬时响应延迟），均已实时捕获并记录，不影响整体业务正常运行。')

    n7 = doc.add_paragraph(style='List Bullet')
    n7.add_run('内容合规监测：').bold = True
    n7.add_run('通过实时巡检机制，对站点所有公开页面文本、核心展示元素、关键信息进行全量扫描核验，重点排查违法违规内容、敏感信息、不合规表述及潜在合规风险点。监测结果显示，站点页面内容整体合规，未发现明确违法违规、敏感及不合规表述，内容安全管控到位，符合平台运营合规要求。')

    doc.add_paragraph('综上，本次监控周期内，站点整体运行状态良好，可用性与内容合规性均达到预期运营标准，未出现影响业务正常开展的重大风险隐患。')




    title11 = doc.add_heading('2. 修复优化建议', level=2)
    title11_run = title11.runs[0]
    title11_run.font.name = '宋体'
    title11_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph("结合本次监控情况，为进一步提升站点运行稳定性、内容合规性，防范潜在风险，提出以下通用性修复及优化建议，适配各类站点长期运营需求：")

    n8 = doc.add_paragraph(style='List Bullet')
    n8.add_run('可用性优化建议：').bold = True

    doc.add_paragraph('针对监控中捕获的瞬时响应延迟、偶尔加载卡顿等轻微异常，建议核查服务器负载、网络带宽及页面资源大小，优化页面加载速度，压缩冗余资源，减少响应耗时。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)

    doc.add_paragraph('建立可用性故障应急修复机制，提前储备常见故障（如无法访问、超时）的修复流程及操作手册，确保一旦出现故障，可快速响应、及时处置，降低故障影响范围。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)

    doc.add_paragraph('定期检查站点访问链路及服务器运行状态，排查潜在硬件、软件故障隐患，及时更新服务器系统及相关组件，保障服务运行环境稳定。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)



    n9 = doc.add_paragraph(style='List Bullet')
    n9.add_run('内容合规优化建议:').bold = True

    doc.add_paragraph('建立内容定期自查机制，结合监控结果，定期对站点历史页面、归档内容进行复盘核查，防范遗漏风险，确保内容合规全覆盖，无死角。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)
    doc.add_paragraph('优化内容发布审核流程，在内容上线前增加合规校验环节，明确审核标准，防范违规内容、敏感信息误上线，从源头保障内容合规。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)
    doc.add_paragraph('及时关注最新合规政策及监管要求，更新内容合规识别标准，同步优化监控系统的合规识别规则，确保监控内容与监管要求保持一致。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)



    n10 = doc.add_paragraph(style='List Bullet')
    n10.add_run('保障建议：').bold = True

    doc.add_paragraph('定期备份站点数据及配置信息，防范数据丢失、配置错乱等问题，确保故障后可快速恢复，降低运营风险。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)

    doc.add_paragraph('安排专人负责监控结果复盘，定期汇总监控数据，分析异常规律，针对性优化监控策略及站点运营管理方案。',style='List Bullet 2').paragraph_format.first_line_indent = Pt(24)



    title12 = doc.add_heading('3. 下一步监测计划', level=2)
    title12_run = title12.runs[0]
    title12_run.font.name = '宋体'
    title12_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


    doc.add_paragraph('为持续保障站点稳定、合规运营，实现风险早发现、早预警、早处置，下一步将延续常态化监控模式，结合本次监控结果及优化建议，完善监控策略，具体计划如下：')


    n11 = doc.add_paragraph(style='List Bullet')
    n11.add_run('1. 优化监控策略，提升监测精准度：').bold = True
    n11.add_run('结合本次监控中的轻微异常及优化建议，调整可用性监测的探测周期（重点时段可适当缩短探测间隔），优化内容合规监测的识别规则，增加高频风险点的扫描频次，提升监控的针对性和精准度，减少误报、漏报情况。')



    n12 = doc.add_paragraph(style='List Bullet')
    n12.add_run('2. 延续核心监测维度，扩大监测覆盖范围：').bold = True
    n12.add_run('持续围绕可用性、内容合规两大核心维度开展监测，同时逐步扩大监测覆盖范围，新增对站点附属页面、关联链路的监测，全面覆盖各类潜在风险点，确保站点全链路、全页面的稳定与合规。')



    n13 = doc.add_paragraph()
    n13.add_run('3. 强化监控数据管理与复盘：').bold = True
    n13.add_run('建立完善的监控数据归档机制，定期汇总监测数据、异常记录及修复情况，每月开展一次监控结果复盘，分析站点运行趋势，排查潜在风险隐患，针对性调整优化建议及监控策略。')



    n14 = doc.add_paragraph()
    n14.add_run('4. 完善预警与处置机制：').bold = True
    n14.add_run('优化监控预警规则，明确不同等级异常的预警方式及处置时限，确保异常情况可及时推送至相关负责人；同步完善故障处置跟踪机制，对出现的异常及修复情况进行全程记录，确保问题闭环解决。')



    n15 = doc.add_paragraph()
    n15.add_run('5. 配合优化落地，跟踪优化效果：').bold = True
    n15.add_run('针对本次提出的修复及优化建议，跟踪优化落地情况，在后续监控过程中重点核查优化效果，确认可用性、内容合规性是否得到进一步提升，及时调整优化方向及监控重点。')


    doc.add_paragraph('下一步，将持续强化监控工作，细化监控流程，完善保障机制，全力支撑站点持续、稳定、合规运营，防范各类可用性及内容合规风险，为业务正常开展提供坚实保障。')














    # 保存文档
    if project_name:
        filename = datetime.datetime.now().strftime(f"{project_name}网站检测服务{report_period}_%Y%m%d_%H%M%S.docx")
    else:
        filename = datetime.datetime.now().strftime(f"网站检测服务{report_period}_%Y%m%d_%H%M%S.docx")

    doc.save(filename)
    return filename


def _format_timedelta(td):
    """将时间差格式化为易读字符串（如 1d 2h 3m）"""
    if td is None:
        return "N/A"
    days, remainder = divmod(td.total_seconds(), 86400)
    hours, remainder = divmod(remainder, 3600)
    minutes, seconds = divmod(remainder, 60)

    parts = []
    if days > 0:
        parts.append(f"{int(days)}d")
    if hours > 0:
        parts.append(f"{int(hours)}h")
    if minutes > 0:
        parts.append(f"{int(minutes)}m")
    if seconds > 0 or not parts:
        parts.append(f"{int(seconds)}s")

    return " ".join(parts)




# --- 主函数 ---
def main():
    print_banner()
    url, username, password,Company, Company_English_name, save_config_needed = handle_credentials()

    try:
        # 连接Uptime Kuma并获取数据
        with UptimeKumaApi(url) as api:
            api.login(username, password)
            print("\n成功连接到Uptime Kuma！")

            if save_config_needed:
                save_config(url, username, Company, Company_English_name)
            
            monitors = api.get_monitors()
    
            if not monitors:
                print("未找到任何监控项，程序退出")
                return

            selected_id = select_monitors(monitors)
            if not selected_id:
                print("未选择任何监控项，程序退出")
                return
            period=chose_report()
            report_times= calculate_hours_since_period_start(period)

            selected_monitors=[]
            all_monitor_data = []
            
            print("\n正在分析数据并生成Word报告...")
            for id in selected_id:
                monitor=api.get_monitor(id)
                selected_monitors.append(monitor)
                


            # 处理每个监控项的数据
            project_name=None
            for monitor in selected_monitors:
                monitor_id = monitor['id']
                monitor_name = monitor['name']
                if monitor.get('parent') is None:
                    if monitor.get("childrenIDs", []) != []:   
                       project_name=monitor_name
                print(f"  - 处理监控项: {monitor_name}")
                

                heartbeats = api.get_monitor_beats(monitor_id, report_times)
                analysis_results = analyze_heartbeats(heartbeats)
             
                summary_stats = calculate_summary_stats(analysis_results)
            

                all_monitor_data.append({
                    "monitor_name": monitor_name,
                    "summary_stats": summary_stats,
                    "downtime_incidents": analysis_results['downtime_incidents'],
                    "keyword_analysis":analysis_results['keyword_analysis'],
                })

            # 生成Word报告
            filename = generate_docx_report(project_name,period,Company, Company_English_name, selected_monitors, all_monitor_data)
            print(f"\n✅ Word报告生成成功: {filename}")

    except UptimeKumaException as e:
        print(f"\n连接Uptime Kuma失败: {e}")
    except Exception as e:
        print(f"\n程序运行出错: {e}")

if __name__ == "__main__":
    main()

